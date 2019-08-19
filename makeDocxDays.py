import datetime, calendar, random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

monthnums = [(i + 7) % 12 +1 for i in range(12)]  # count from august to july
mos=',janvier,février,mars,avril,mai,juin,juillet,août,septembre,octobre,novembre,décembre'.split(',')
dow = 'LUNDI,MARDI,MERCREDI,JEUDI,VENDREDI,SAMEDI,DIMANCHE'.split(',')


# This is a kludge - dates for 2017-2018 given as dictionary!!!
schooldates = {'2018-01-31': 'Jour 5', '2017-08-29': 'Journée pédagogique', '2017-08-28': 'Journée pédagogique', '2018-06-11': 'Jour 4', '2018-06-13': 'Jour 6', '2018-06-12': 'Jour 5', '2017-12-19': 'Jour 3', '2017-12-18': 'Jour 2', '2018-06-08': 'Jour 3', '2017-10-09': 'congé ferié', '2017-10-06': 'Jour 6', '2017-12-12': 'Jour 4', '2017-10-04': 'Jour 4', '2017-10-05': 'Jour 5', '2017-11-30': 'Jour 3', '2017-10-03': 'Jour 3', '2017-12-15': 'Jour 1', '2017-12-14': 'Jour 6', '2017-12-22': 'Jour 6', '2017-12-20': 'Jour 4', '2017-12-21': 'Jour 5', '2017-09-08': 'Jour 5', '2017-09-04': 'congé ferié', '2017-09-05': 'Jour 2', '2017-09-06': 'Jour 3', '2017-09-07': 'Jour 4', '2017-09-01': 'Jour 1', '2017-12-27': 'congé', '2018-06-06': 'Jour 1', '2017-10-19': 'Jour 2', '2017-10-18': 'Jour 1', '2017-10-17': 'Jour 6', '2017-10-16': 'Jour 5', '2017-10-11': 'Jour 2', '2017-10-10': 'Jour 1', '2017-10-13': 'Jour 4', '2017-10-12': 'Jour 3', '2017-11-23': 'Journée pédagogique', '2017-11-22': 'Jour 5', '2017-11-21': 'Jour 4', '2017-11-20': 'Jour 3', '2017-11-27': 'Jour 6', '2017-11-24': 'Journée pédagogique', '2017-11-29': 'Jour 2', '2017-11-28': 'Jour 1', '2018-04-11': 'Jour 5', '2018-06-07': 'Jour 2', '2017-12-29': 'congé', '2018-02-23': 'Journée pédagogique', '2018-03-19': 'Jour 6', '2018-03-16': 'Jour 5', '2018-03-14': 'Jour 3', '2018-03-15': 'Jour 4', '2018-03-12': 'Jour 1', '2018-03-13': 'Jour 2', '2017-09-13': 'Jour 2', '2018-04-23': 'Jour 1', '2018-04-20': 'Jour 6', '2018-04-26': 'Jour 4', '2018-04-27': 'Jour 5', '2017-09-15': 'Jour 4', '2017-09-14': 'Jour 3', '2018-01-30': 'Jour 4', '2017-09-18': 'Jour 5', '2017-10-20': 'Jour 3', '2017-10-23': 'Journée pédagogique', '2017-10-24': 'Jour 4', '2017-10-25': 'Jour 5', '2017-10-26': 'Jour 6', '2017-10-27': 'Jour 1', '2018-06-05': 'Jour 6', '2018-06-18': 'Jour 3', '2018-04-04': 'congé', '2018-06-27': 'Journée pédagogique', '2018-03-09': 'congé', '2018-03-08': 'congé', '2018-03-05': 'congé', '2018-03-07': 'congé', '2018-03-06': 'congé', '2018-03-01': 'Jour 5', '2018-03-02': 'Jour 6', '2018-01-29': 'Jour 3', '2018-04-30': 'Jour 6', '2018-01-23': 'Jour 6', '2018-01-22': 'Jour 5', '2018-01-25': 'Jour 2', '2018-01-24': 'Jour 1', '2018-01-26': 'Journée pédagogique', '2017-12-13': 'Jour 5', '2017-10-31': 'Jour 3', '2017-10-30': 'Jour 2', '2017-12-11': 'Jour 3', '2017-12-05': 'Jour 6', '2017-10-02': 'Jour 2', '2018-04-09': 'Jour 3', '2018-05-18': 'Journée pédagogique', '2018-03-30': 'congé', '2018-04-05': 'congé', '2018-04-06': 'congé', '2018-01-18': 'Jour 3', '2018-01-19': 'Jour 4', '2018-04-02': 'congé', '2018-04-03': 'congé', '2018-01-15': 'Jour 6', '2018-01-16': 'Jour 1', '2018-01-17': 'Jour 2', '2018-01-10': 'Jour 3', '2018-01-11': 'Jour 4', '2018-01-12': 'Jour 5', '2018-05-25': 'Jour 5', '2018-05-24': 'Jour 4', '2018-06-01': 'Jour 4', '2018-05-21': 'congé ferié', '2018-05-23': 'Jour 3', '2018-05-22': 'Jour 2', '2018-05-29': 'Jour 1', '2018-05-28': 'Jour 6', '2018-02-02': 'Jour 0 - Examens', '2017-12-04': 'Jour 5', '2018-02-01': 'Jour 6', '2018-02-06': 'Jour 1', '2018-02-07': 'Jour 2', '2018-02-05': 'Jour 0 - Examens', '2018-02-08': 'Jour 3', '2018-02-09': 'Jour 4', '2018-03-27': 'Jour 6', '2018-03-26': 'Jour 5', '2018-04-19': 'Jour 5', '2018-04-18': 'Jour 4', '2017-12-26': 'congé', '2018-03-22': 'Jour 3', '2018-03-21': 'Jour 2', '2017-12-25': 'congé', '2018-04-13': 'Jour 1', '2018-04-12': 'Jour 6', '2017-12-28': 'congé', '2018-04-10': 'Jour 4', '2018-04-17': 'Jour 3', '2018-04-16': 'Jour 2', '2018-03-29': 'Jour 2', '2018-03-28': 'Jour 1', '2018-01-03': 'congé', '2018-01-02': 'congé', '2018-01-01': 'congé', '2018-01-05': 'congé', '2018-01-04': 'congé', '2018-05-10': 'Jour 2', '2018-05-11': 'Jour 3', '2018-01-09': 'Jour 2', '2018-01-08': 'Jour 1', '2018-05-14': 'Jour 4', '2018-05-15': 'Jour 5', '2018-05-16': 'Jour 6', '2018-05-17': 'Jour 1', '2018-02-13': 'Jour 6', '2018-02-12': 'Jour 5', '2018-02-15': 'Jour 2', '2018-02-14': 'Jour 1', '2018-02-16': 'Jour 3', '2018-02-19': 'Jour 4', '2018-03-23': 'Jour 4', '2018-06-04': 'Jour 5', '2018-06-19': 'Jour 0 - Examens', '2018-04-25': 'Jour 3', '2018-05-30': 'Jour 2', '2018-06-15': 'Jour 2', '2017-11-13': 'Jour 4', '2017-11-10': 'Jour 3', '2017-11-16': 'Jour 1', '2017-11-17': 'Jour 2', '2017-11-14': 'Jour 5', '2017-11-15': 'Jour 6', '2018-05-09': 'Jour 1', '2018-05-08': 'Jour 6', '2018-05-07': 'Jour 5', '2018-04-24': 'Jour 2', '2018-05-04': 'Jour 4', '2018-05-03': 'Jour 3', '2018-05-02': 'Jour 2', '2018-05-01': 'Jour 1', '2018-02-26': 'Jour 2', '2018-02-27': 'Jour 3', '2018-02-20': 'Jour 5', '2018-02-21': 'Jour 6', '2018-02-22': 'Jour 1', '2017-09-12': 'Jour 1', '2018-05-31': 'Jour 3', '2018-02-28': 'Jour 4', '2017-09-11': 'Jour 6', '2017-08-30': 'Accueil des élèves', '2017-08-31': 'Accueil des élèves', '2017-09-28': 'Jour 6', '2017-09-29': 'Jour 1', '2017-09-22': 'Journée pédagogique', '2017-09-20': 'Jour 1', '2017-09-21': 'Jour 2', '2017-09-26': 'Jour 4', '2017-09-27': 'Jour 5', '2017-09-25': 'Jour 3', '2017-11-09': 'Jour 2', '2017-11-08': 'Jour 1', '2017-12-06': 'Jour 1', '2017-12-07': 'Jour 2', '2017-12-01': 'Jour 4', '2017-11-01': 'Jour 4', '2017-11-03': 'Jour 0 - Examens', '2017-11-02': 'Jour 5', '2017-12-08': 'Journée pédagogique', '2017-11-07': 'Jour 6', '2017-11-06': 'Jour 0 - Examens', '2018-06-20': 'Jour 0 - Examens', '2018-06-21': 'Jour 0 - Examens', '2018-06-22': 'Journée pédagogique', '2018-03-20': 'Jour 1', '2018-06-25': 'congé ferié', '2018-06-26': 'Journée pédagogique', '2018-06-14': 'Jour 1', '2018-06-28': 'Journée pédagogique', '2017-09-19': 'Jour 6'}

# strategy:
#  go through months of schoolyear
#   for each month, create a table
#     add_row().cells[0].text = upper monthname + text
#     add_row() for day names
#     for each week in calendar.monthcalendar, add_row for day info
#               if 0, skip
#               if not present, just put date
#               put date and Jour

doc = Document()

font = doc.styles['Normal'].font
font.name = 'Calibri'
font.size = Pt(9)


run = doc.add_heading('Évaluations')

section = doc.sections[-1]
half_inch = 914400 // 2
section.left_margin = half_inch
section.right_margin = half_inch

for thismonth in monthnums:
    monthname = mos[thismonth]
    monthnumstr = '{:02d}'.format(thismonth)
    dates = { d:j for d,j in schooldates.items() if d[5:7] == monthnumstr }
    if len(dates) == 0: break
    year = list(dates.keys())[0][:4]
    table = doc.add_table(rows=1,cols=7)
    hdr = table.rows[0].cells
    hdrtext = hdr[0].add_paragraph('{} {} – CALENDRIER DES ÉVALUATIONS'.format(monthname.upper(), year))
    hdrtext.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr[0].merge(hdr[6])
    #hdr[0].text = '{} {} – CALENDRIER DES ÉVALUATIONS'.format(monthname.upper(), year)
    dayline = table.add_row().cells
    for n,d in enumerate(dow):
        dtext = dayline[n].text = d
#        dtext.alignment = WD_ALIGN_PARAGRAPH.CENTER
    weeks = calendar.monthcalendar(int(year), thismonth)
    for week in weeks:
        daterow = table.add_row().cells
        for n,day in enumerate(week):
            if day != 0:
                dkey = '{}-{:02d}-{:02d}'.format(year,thismonth,day)
                info = schooldates.get(dkey) or ''
                info = ' - ' + info if info else ''
                dtext = '{} {}{}'.format(day, monthname, info)
                daterow[n].text = dtext
        table.add_row()

    tab_style = 'Medium Shading 1 Accent {:d}'.format(random.randint(1,6))
    table.style = doc.styles[tab_style]
    doc.add_page_break()
    
doc.save('Évaluations.docx')
